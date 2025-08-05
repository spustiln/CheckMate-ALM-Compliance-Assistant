import streamlit as st
import pandas as pd
from datetime import date
from docx import Document
from io import BytesIO
import tempfile
import os

st.set_page_config(page_title="CheckMate AML Assistant")
st.title("♟️ CheckMate AML Assistant")

# Role selection
role = st.sidebar.selectbox("Select Role", ["Tier 1 Analyst", "Tier 2 Compliance Officer"])

# File upload
st.sidebar.header("Upload Files")
tx_file = st.sidebar.file_uploader("Upload transactions CSV", type="csv")
alerts_file = st.sidebar.file_uploader("Upload alerts CSV", type="csv")

if tx_file:
    tx = pd.read_csv(tx_file)

    if alerts_file:
        alerts = pd.read_csv(alerts_file)
        tx["alert_flag"] = tx["ACCOUNT_ID"].isin(alerts["ACCOUNT_ID"])
    else:
        tx["alert_flag"] = False

    amount_col = next((c for c in tx.columns if "AMOUNT" in c.upper()), None)
    if not amount_col:
        st.error("Missing amount column.")
        st.stop()

    if "TXN_STEP" in tx.columns:
        tx["Case_ID"] = tx["TXN_STEP"].astype(str) + "-" + tx["TXN_ID"].astype(str)
    else:
        tx["Case_ID"] = tx["TXN_ID"].astype(str)

    if "COUNTER_PARTY_ACCOUNT_NUM" not in tx.columns:
        tx["COUNTER_PARTY_ACCOUNT_NUM"] = ""

    tx["Flags"] = tx.apply(lambda r: [
        reason for reason in [
            "High-value" if r[amount_col] > 150 else None,
            "Prior alert" if r.get("alert_flag") else None,
            "Frequent receiver" if tx["COUNTER_PARTY_ACCOUNT_NUM"].value_counts().get(r["COUNTER_PARTY_ACCOUNT_NUM"], 0) > 2 else None
        ] if reason is not None
    ], axis=1)

    flagged = tx[tx["Flags"].map(len) > 0].copy()
    flagged["Flags_str"] = flagged["Flags"].map("; ".join)

    if role == "Tier 1 Analyst":
        st.header("🧑‍💼 Tier 1 Review")
        st.dataframe(flagged[["Case_ID", "ACCOUNT_ID", amount_col, "Flags_str"]])

        sel_case = st.selectbox("Select Case ID", ["" ] + flagged["Case_ID"].tolist())

        if sel_case:
            if 'last_case' not in st.session_state:
                st.session_state['last_case'] = ""

            if sel_case != st.session_state['last_case']:
                st.session_state['tier1_decision'] = None
                st.session_state['tier1_actions'] = []
                st.session_state['tier2_approval'] = None
                st.session_state['last_case'] = sel_case

            row = flagged[flagged["Case_ID"] == sel_case].iloc[0]
            st.write("**Flags:**", row["Flags_str"])
            suspicious = st.radio("Tier 1 Analyst - Suspicious?", ["No", "Yes"], key="tier1_decision")

            tier1_actions = st.multiselect("Tier 1 Actions", [
                "Stop the transaction in question",
                "Place temporary hold on sender account",
                "Contact Account Owner for clarification",
                "Monitor Only",
                "Escalate to Tier 2"
            ], key="tier1_actions")

            if "cases" not in st.session_state:
                st.session_state.cases = []

            if st.button("Submit Tier 1 Review"):
                st.session_state.cases.append({
                    "Case_ID": row["Case_ID"],
                    "Sender": row["ACCOUNT_ID"],
                    "Receiver": row["COUNTER_PARTY_ACCOUNT_NUM"],
                    "Amount": row[amount_col],
                    "Reason": row["Flags_str"],
                    "Action": ", ".join(tier1_actions)
                })
                st.success("Submitted for Tier 2 Review")

    elif role == "Tier 2 Compliance Officer":
        st.header("🧑‍⚖️ Tier 2 Review Queue")
        if "cases" not in st.session_state or not st.session_state.cases:
            st.info("No cases escalated yet.")
        else:
            df = pd.DataFrame(st.session_state.cases)
            escalated = df[df["Action"].str.contains("Escalate", na=False)]

            if escalated.empty:
                st.info("No cases awaiting Tier 2 review.")
            else:
                choice = st.selectbox("Review escalated case:", escalated["Case_ID"])
                selected = escalated[escalated["Case_ID"] == choice].iloc[0]

                st.write(f"**Sender:** {selected['Sender']}")
                st.write(f"**Amount:** ${selected['Amount']}")
                st.write(f"**Flag Reason:** {selected.get('Reason', 'N/A')}")

                approve = st.radio("Approve SAR Filing?", ["No", "Yes"], key="tier2_final_approval")

                if st.button("Generate SAR (Tier 2)"):
                    doc = Document()
                    doc.add_heading(f"SAR – Case {selected['Case_ID']}", 0)
                    doc.add_paragraph(f"Date: {date.today()}")
                    doc.add_paragraph(f"Sender: {selected['Sender']}")
                    doc.add_paragraph(f"Amount: ${selected['Amount']}")
                    doc.add_paragraph("Tier 2 Approval: " + approve)

                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                        temp_path = tmp.name
                        doc.save(temp_path)
                        with open(temp_path, "rb") as f:
                            st.download_button("Download SAR Report", f, file_name=f"SAR_{selected['Case_ID']}.docx")
                    os.unlink(temp_path)

