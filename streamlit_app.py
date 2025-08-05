import streamlit as st
import pandas as pd
from datetime import date, timedelta
from io import BytesIO
from docx import Document

# Configuration
BSA_THRESHOLD = 9000
REVIEW_INTERVAL = timedelta(days=90)

# Helper Functions
def ai_flag_reasons(row, amount_col, tx_df):
    labels = []
    if row[amount_col] > BSA_THRESHOLD:
        labels.append(f"High-value: ${row[amount_col]:,.2f}")
    mask = (tx_df["ACCOUNT_ID"] == row["ACCOUNT_ID"]) & (tx_df.get("COUNTER_PARTY_ACCOUNT_NUM") == row.get("COUNTER_PARTY_ACCOUNT_NUM"))
    if tx_df[mask].shape[0] > 3:
        labels.append("Frequent receiver pattern (possible mule)")
    if row.get("alert_flag"):
        labels.append("Previous alert on sender")
    if not labels:
        labels.append("Threshold exceeded")
    return labels

def generate_sar_docx(case_id, row, flags, tier1_actions, tier2_approved, amount_col):
    doc = Document()
    today = date.today().strftime("%Y-%m-%d")

    doc.add_heading(f"SUSPICIOUS ACTIVITY REPORT — Case {case_id}", level=1)
    doc.add_paragraph(f"Report Date: {today}")
    doc.add_paragraph("Institution: Simulated Bank Corp")
    doc.add_paragraph("Filed by: AI Compliance Simulator")

    doc.add_heading("Section I – Subject Information", level=2)
    fields = {
        "Sender Account ID": row["ACCOUNT_ID"],
        "Receiver Account ID": row.get("COUNTER_PARTY_ACCOUNT_NUM", "N/A"),
        "Amount": f"${row[amount_col]:,.2f}",
        "Channel": row.get("TXN_SOURCE_TYPE_CODE", "N/A"),
        "Step": row.get("TXN_STEP", "N/A"),
        "Previous Alert": "Yes" if row.get("alert_flag") else "No"
    }
    for label, val in fields.items():
        doc.add_paragraph(f"{label}: {val}")

    doc.add_heading("Section II – Suspicious Activity Indicators", level=2)
    for f in flags:
        doc.add_paragraph(f"• {f}")

    doc.add_heading("Section III – Narrative Explanation", level=2)
    doc.add_paragraph(f"Transfer of {fields['Amount']} via {fields['Channel']} at Step {fields['Step']}.")
    doc.add_paragraph(f"AI Flags: {', '.join(flags)}.")
    if tier1_actions:
        doc.add_paragraph(f"Tier 1 Actions: {', '.join(tier1_actions)}.")
    doc.add_paragraph("Tier 2 Approval: " + ("Yes" if tier2_approved else "No"))

    doc.add_heading("Section IV – Actions Taken", level=2)
    doc.add_paragraph("Analyst Decision: " + ("Yes" if "Escalate" in tier1_actions else "No"))
    doc.add_paragraph(f"Tier 1: {', '.join(tier1_actions)}")
    doc.add_paragraph(f"Tier 2 Approved: {'Yes' if tier2_approved else 'No'}")
    return doc

# Streamlit UI
st.set_page_config(page_title="AML Simulator", layout="wide")
st.title("💼 AML Compliance Simulator")

tx_file = st.sidebar.file_uploader("Upload transactions CSV", type="csv")
alerts_file = st.sidebar.file_uploader("Upload alerts CSV", type="csv")

if tx_file:
    tx = pd.read_csv(tx_file)
    if alerts_file:
        alerts = pd.read_csv(alerts_file)
        tx["alert_flag"] = tx["ACCOUNT_ID"].isin(alerts["ACCOUNT_ID"])
    else:
        tx["alert_flag"] = False

    amount_col = next((c for c in tx.columns if "AMOUNT" in c.upper()), None)
    if not amount_col:
        st.error("Missing amount column.")
        st.stop()

    if "COUNTER_PARTY_ACCOUNT_NUM" not in tx.columns:
        tx["COUNTER_PARTY_ACCOUNT_NUM"] = ""

    tx["Case_ID"] = tx.get("TXN_STEP", 0).astype(str) + "-" + tx["TXN_ID"].astype(str)
    tx["Flags"] = tx.apply(lambda r: ai_flag_reasons(r, amount_col, tx), axis=1)
    flagged = tx[tx["Flags"].map(len) > 0].copy()
    flagged["Flags_str"] = flagged["Flags"].map("; ".join)

    st.header("Flagged Transactions")
    st.dataframe(flagged[["Case_ID", "ACCOUNT_ID", amount_col, "Flags_str"]])

    sel_case = st.selectbox("Select Case ID", [""] + flagged["Case_ID"].tolist())
    if sel_case:
        row = flagged[flagged["Case_ID"] == sel_case].iloc[0]
        st.subheader("Review Case")
        st.write("**Flags:**", row["Flags_str"])
        suggestion = "Escalate to Tier 2" if "High-value" in row["Flags_str"] else "Monitor"
        st.write("Suggested Action:", suggestion)

        suspicious = st.radio("Suspicious?", ["No", "Yes"], index=1)
        tier1_actions = []
        if suspicious == "Yes":
            tier1_actions = st.multiselect("Tier 1 Actions", ["Block transaction", "Place hold", "Escalate to Tier 2"])
            if not tier1_actions:
                tier1_actions = ["Monitor"]

        tier2_approved = False
        if "Escalate to Tier 2" in tier1_actions:
            tier2_approved = st.radio("Tier 2 Approval?", ["No", "Yes"]) == "Yes"

        if st.button("📄 Generate SAR"):
            doc = generate_sar_docx(sel_case, row, row["Flags"], tier1_actions, tier2_approved, amount_col)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("Download SAR (DOCX)", buffer, file_name=f"SAR_{sel_case}.docx")
